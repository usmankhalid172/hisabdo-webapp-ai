from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import zipfile, os, math

OUT = Path('/home/user/HisabDo_AI_Model_API_Research')
OUT.mkdir(exist_ok=True)
DATE = '20 August 2026'

SOURCES = [
('S1','OpenAI — Models overview (GPT-5.6 family, context, tools, prices)','https://platform.openai.com/docs/models'),
('S2','OpenAI — Structured model outputs','https://platform.openai.com/docs/guides/structured-outputs'),
('S3','OpenAI — API data controls and retention','https://platform.openai.com/docs/guides/your-data'),
('S4','OpenAI — text-embedding-3-small','https://platform.openai.com/docs/models/text-embedding-3-small'),
('S5','OpenAI — text-embedding-3-large','https://platform.openai.com/docs/models/text-embedding-3-large'),
('S6','Google — Gemini Developer API pricing','https://ai.google.dev/gemini-api/docs/pricing'),
('S7','Google — Gemini 3.1 Flash-Lite model page','https://ai.google.dev/gemini-api/docs/models/gemini-3.1-flash-lite'),
('S8','Google — Gemini 3.7 Flash guide','https://ai.google.dev/gemini-api/docs/latest-model'),
('S9','Google — Gemini Embeddings documentation','https://ai.google.dev/gemini-api/docs/embeddings'),
('S10','Google — Gemini API Additional Terms (data use)','https://ai.google.dev/gemini-api/terms'),
('S11','Google — Gemini Developer API zero data retention','https://ai.google.dev/gemini-api/docs/zdr'),
('S12','Anthropic — Claude models overview','https://docs.anthropic.com/en/docs/about-claude/models/overview'),
('S13','Anthropic — Claude pricing','https://docs.anthropic.com/en/docs/about-claude/pricing'),
('S14','Anthropic — Claude release notes','https://docs.anthropic.com/en/release-notes/overview'),
('S15','Anthropic — API and data retention','https://docs.anthropic.com/en/docs/build-with-claude/zero-data-retention'),
('S16','Mistral — Model selection guide','https://docs.mistral.ai/models/model-selection-guide'),
('S17','Mistral — API pricing','https://docs.mistral.ai/inference/pricing'),
('S18','Mistral — Known limitations (JSON mode)','https://docs.mistral.ai/resources/known-limitations'),
('S19','Qwen — Function calling documentation','https://qwen.readthedocs.io/en/latest/framework/function_call.html'),
('S20','Qwen — Qwen3 quickstart and OpenAI-compatible serving','https://qwen.readthedocs.io/en/latest/getting_started/quickstart.html'),
('S21','Meta — Llama models repository/model index','https://github.com/meta-llama/llama-models'),
('S22','Hugging Face — Inference Providers pricing and credits','https://huggingface.co/docs/inference-providers/pricing'),
('S23','Hugging Face — Inference Endpoints pricing','https://huggingface.co/docs/inference-endpoints/pricing'),
('S24','Sentence Transformers — Documentation','https://sbert.net/'),
('S25','Sentence Transformers — Inference efficiency/ONNX/OpenVINO','https://sbert.net/docs/sentence_transformer/usage/efficiency.html'),
('S26','BAAI — BGE-M3 model card','https://huggingface.co/BAAI/bge-m3'),
('S27','intfloat — multilingual-e5-small model card','https://huggingface.co/intfloat/multilingual-e5-small'),
('S28','Voyage AI — Embedding model choices','https://docs.voyageai.com/docs/embeddings'),
('S29','Voyage AI — Pricing','https://docs.voyageai.com/docs/pricing'),
('S30','Qdrant — Managed Cloud and free-tier details','https://qdrant.tech/cloud/'),
('S31','FastAPI — Handling errors','https://fastapi.tiangolo.com/tutorial/handling-errors/'),
('S32','FastAPI — Server-Sent Events','https://fastapi.tiangolo.com/tutorial/server-sent-events/'),
('S33','FastAPI — OAuth2 scopes','https://fastapi.tiangolo.com/advanced/security/oauth2-scopes/'),
('S34','Pydantic — TypeAdapter validation and JSON schema','https://docs.pydantic.dev/latest/api/type_adapter/'),
('S35','scikit-learn — Working with text data and evaluation','https://scikit-learn.org/stable/tutorial/text_analytics/working_with_text_data.html'),
('S36','scikit-learn — F1 score','https://scikit-learn.org/stable/modules/generated/sklearn.metrics.f1_score.html'),
]

NAVY='17324D'; BLUE='1F6E8C'; TEAL='159A9C'; LIGHT='EAF3F5'; PALE='F6F8FA'; DARK='1F2933'; RED='A63D40'; GOLD='B7791F'; GREEN='2F855A'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell_text_color(cell, color='FFFFFF', bold=False, size=8):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string(color); r.bold=bold; r.font.size=Pt(size)
def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def base_doc(title, subtitle):
    d=Document(); sec=d.sections[0]
    sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
    styles=d.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9); styles['Normal'].font.color.rgb=RGBColor.from_string(DARK)
    styles['Normal'].paragraph_format.space_after=Pt(4)
    for name,size,color in [('Title',28,NAVY),('Heading 1',18,NAVY),('Heading 2',13,BLUE),('Heading 3',10.5,TEAL)]:
        st=styles[name]; st.font.name='Aptos Display'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
        st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(5)
    if 'Callout' not in styles:
        st=styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH); st.font.name='Aptos'; st.font.size=Pt(9); st.font.color.rgb=RGBColor.from_string(NAVY); st.font.bold=True
        st.paragraph_format.left_indent=Inches(.18); st.paragraph_format.right_indent=Inches(.18); st.paragraph_format.space_before=Pt(5); st.paragraph_format.space_after=Pt(5)
    # cover
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55)
    r=p.add_run('HISABDO WEB APP AI'); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(TEAL)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8)
    r=p.add_run(title); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle); r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
    r=p.add_run(f'Research cut-off: {DATE}\nPrepared for technical team / supervisor submission'); r.font.size=Pt(10)
    t=d.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    vals=[('SCOPE','LLM • RAG • categorization'),('DECISION','Capstone-first, production-aware'),('EVIDENCE','Official sources preferred')]
    for c,(a,b) in zip(t.rows[0].cells,vals):
        shade(c,LIGHT); c.text=a+'\n'+b; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c,120,120,120,120)
        for i,r in enumerate(c.paragraphs[0].runs): r.font.color.rgb=RGBColor.from_string(NAVY); r.bold=(i==0)
    d.add_page_break()
    # header/footer
    header=sec.header.paragraphs[0]; header.text='HisabDo AI — Technical Research'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(BLUE)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=footer.add_run('HisabDo • Research cut-off '+DATE+' • ')
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)
    for r in footer.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(BLUE)
    return d

def heading(d,text,level=1): d.add_heading(text,level=level)
def para(d,text='',bold_prefix=None,style=None):
    p=d.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p
def bullets(d,items,level=0):
    for x in items:
        p=d.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(x)
def numbered(d,items):
    for x in items: d.add_paragraph(x,style='List Number')
def callout(d,title,text,color=LIGHT):
    t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,color); set_cell_margins(c,130,160,130,160)
    p=c.paragraphs[0]; r=p.add_run(title+' — '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    d.add_paragraph().paragraph_format.space_after=Pt(0)
def table(d,headers,rows,widths=None,font=7.5):
    t=d.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c,NAVY); set_cell_text_color(c,'FFFFFF',True,font); set_cell_margins(c)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ridx%2: shade(cells[i],PALE)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs: r.font.size=Pt(font)
    t.autofit=True
    return t

def refs(d, ids=None):
    heading(d,'References',1)
    wanted=set(ids) if ids else None
    for sid,title,url in SOURCES:
        if wanted is None or sid in wanted:
            p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(.15); p.paragraph_format.first_line_indent=Inches(-.15)
            p.add_run(f'[{sid}] {title}. ').bold=True; p.add_run(url)
    para(d,f'All web sources were checked for this report on {DATE}. Prices, model aliases, previews, quotas, regional availability and retention terms can change; verify official pages immediately before implementation or procurement.',style='Callout')

def add_toc_note(d):
    heading(d,'Document Map',1)
    table(d,['Section','Purpose'],[
        ('1–2','Executive decision and HisabDo requirements'),('3','Major LLM/API comparison'),('4','RAG and embedding decision'),('5','Expense categorization decision'),('6','FastAPI architecture and contracts'),('7–10','Recommendations, cost, risk, stack and team actions'),('Appendices','Evaluation gates, decision record and references')
    ],font=8.5)
    para(d,'Note: Microsoft Word can generate a live table of contents from the built-in Heading styles (References → Table of Contents).')

# ---------------- Main report ----------------
d=base_doc('AI Models & APIs — Technical Research Report','Practical recommendations for the financial assistant, RAG, expense categorization and FastAPI service layer')
add_toc_note(d)
heading(d,'Decision Status and Evidence Rules',1)
callout(d,'Primary recommendation','Use GPT-5.6 Luna behind a provider-neutral adapter for the assistant; local BGE-M3 + Qdrant for RAG; and a rules + TF-IDF Logistic Regression expense classifier. The LLM must explain, not perform authoritative ledger calculations.')
bullets(d,[
'“Verified” statements are tied to official documentation citations [S#]. “Recommendation” statements are engineering judgments for HisabDo and must be validated on a project-specific test set.',
'No public model specification proves financial correctness. The selected default is a cost/reliability/implementation trade-off—not a claim of universal superiority.',
'Pricing is quoted only where an official page exposed a current rate on the research date. Taxes, currency conversion, cached tokens, reasoning/thinking tokens, search/tool fees, hosting, egress and support are not silently included.',
'This report assumes a capstone: a small team, limited budget, modest knowledge base, and a need to demonstrate sound engineering rather than maximize benchmark scores.'
])

heading(d,'1. Executive Summary',1)
para(d,'HisabDo needs three different forms of intelligence and should not solve all three with one generative model. The financial assistant needs controlled natural-language generation and tool calling; RAG needs high-quality multilingual retrieval and citations; expense categorization needs repeatable, inexpensive classification with measurable confidence. Treating these as separate services lowers cost and makes failures easier to test.')
para(d,'For the conversational assistant, GPT-5.6 Luna is the recommended capstone default because the current OpenAI model catalogue lists a 1.05M-token context, function tools and a low published token price ($0.20/M input and $1.20/M output), while OpenAI documents schema-adherent Structured Outputs [S1–S2]. This is a provisional engineering selection: HisabDo should run a 100–200 question financial/RAG evaluation against Gemini 3.1 Flash-Lite and retain the winner only if it meets safety and quality gates.')
para(d,'For RAG, run BAAI/bge-m3 locally and store 1024-dimensional vectors in Qdrant. BGE-M3 supports more than 100 languages, up to 8,192 tokens, and dense/sparse/multi-vector retrieval [S26]; that is useful for English, Urdu and Roman-Urdu experiments and avoids embedding sensitive knowledge through a third party. Use hybrid lexical + dense retrieval and citations. If local resources are tight, multilingual-e5-small (384 dimensions) is the lightweight alternative [S27].')
para(d,'For expenses, use deterministic merchant/keyword rules followed by word- and character-ngram TF-IDF with multinomial Logistic Regression. This gives fast CPU inference, inspectable errors and calibrated probabilities. Route low-confidence predictions to “Needs review,” learn from user corrections, and use an LLM only as an optional anonymized fallback—not as the primary classifier.')
para(d,'FastAPI should expose versioned endpoints, validate all requests and model outputs with Pydantic, isolate providers behind adapters, apply timeouts/circuit breakers, and return safe degraded responses. Financial facts (balances, totals, budgets, due dates) must come from authenticated application tools/database queries; the LLM may phrase the result but must not invent or recompute authoritative values.')

heading(d,'2. Requirements Analysis',1)
heading(d,'2.1 Functional needs',2)
table(d,['Workstream','Actual need','Do not delegate blindly'],[
('Assistant','Understand intent; explain budgets/transactions; call read-only tools; ask clarifying questions; return a typed response.','Ledger arithmetic, transaction mutation, tax/investment decisions, or claims unsupported by tools/RAG.'),
('RAG','Retrieve HisabDo policies/help, category definitions and financial education; cite exact chunks; abstain when evidence is weak.','Treat model pretraining as current HisabDo policy or fabricate citations.'),
('Categorization','Map noisy merchant/description/amount/channel fields to a controlled taxonomy; return confidence; accept correction.','Free-form category invention or high-cost LLM calls for every transaction.'),
('Service layer','Stable API contracts, authentication, observability, throttling, fallbacks and versioned models.','Expose provider keys, raw prompts, hidden reasoning, stack traces or unbounded user content.')
],font=7.8)
heading(d,'2.2 Non-functional priorities',2)
bullets(d,[
'Reliability before fluency: cite sources, validate schemas, use deterministic tools, and provide abstention.',
'Privacy by minimization: redact names/account numbers; send only fields necessary for the task; use paid/commercial data terms for real financial records.',
'Capstone practicality: one Python service, CPU-friendly classifier, local vector store, a single primary LLM API, and a thin provider adapter.',
'Measurability: macro-F1 and selective accuracy for categories; Recall@k/MRR for retrieval; groundedness, citation precision, refusal quality, latency and cost per successful answer for the assistant.',
'Change tolerance: pin model IDs where available, place aliases/config in environment variables, maintain contract tests, and document deprecation checks.'
])
heading(d,'2.3 Minimum safety behavior',2)
numbered(d,[
'Classify the intent and risk. General education and app guidance are allowed; high-stakes personalized tax, legal, lending or investment recommendations must be limited, caveated or escalated.',
'Retrieve approved knowledge and/or call a deterministic read-only financial tool.',
'Reject instructions found inside retrieved documents that attempt to override system rules (prompt injection).',
'Generate a structured draft with citations and uncertainty flags.',
'Validate fields, cited chunk IDs, numeric consistency and policy rules. If validation fails, retry once with a repair prompt, then return a safe fallback.',
'Log redacted telemetry and collect explicit user feedback without storing secrets or full financial content by default.'
])

heading(d,'3. LLM/API Comparison',1)
para(d,'Prices below are USD per one million text tokens where an official page provided a rate on the research date. “Fit” is a HisabDo engineering judgment. Preview models are not recommended as the only production dependency.')
llm_rows=[
('GPT-5.6 Luna','OpenAI','Hosted proprietary','$0.20 in / $1.20 out; 1.05M context; functions/search/file-search; schema outputs [S1–S2]','Closed; API dependency; financial quality still requires tests','Yes','Yes','Yes','Primary: high'),
('GPT-5.6 Terra','OpenAI','Hosted proprietary','$2 in / $12 out; same large context/tool surface; stronger tier [S1]','12× Luna output price; overkill for routine chat','Yes','Yes','Yes','Escalation only'),
('Gemini 3.1 Flash-Lite','Google','Hosted proprietary','Stable; 1,048,576 input/65,536 output; function calling, structured outputs, search/file search [S7]','Free tier content is used to improve products; use paid tier for data [S6,S10]','Yes','Yes','Yes','Secondary: high'),
('Gemini 3.7 Flash','Google','Hosted proprietary','1M context; tunable thinking; current promotional standard pricing $0.75/$3.75 to 31 Dec 2026 [S8]','Price changes in 2027; newest model needs project testing','Yes','Yes','Yes','Quality challenger'),
('Claude Sonnet 5','Anthropic','Hosted proprietary','$2/$10; 1M context, 128k output; fast balanced tier [S12–S14]','Higher routine cost; separate SDK/contracts','Yes','Yes','Yes','Premium secondary'),
('Claude Haiku 4.5','Anthropic','Hosted proprietary','$1/$5; 200k context; fastest Claude tier [S12–S13]','Less economical than lowest-cost alternatives for this capstone','Yes','Yes','Yes','Optional fallback'),
('Mistral Small 4 API','Mistral','Hosted/open-weight family','$0.15/$0.60 on current guide; function calling and structured outputs [S16–S17]','Must verify exact model/version and hosted feature parity','Yes','Yes','Yes','Strong low-cost alternative'),
('Qwen3 8B / newer Qwen open weights','Qwen ecosystem','Self-host/open-weight','Local control; vLLM/SGLang OpenAI-compatible serving; Hermes-style tools [S19–S20]','GPU/RAM/operations; tool calls are template/parser dependent; lower assurance','Yes','Conditional','Prompt/schema dependent','Offline/degraded fallback'),
('Llama 4 Scout/Maverick','Meta/hosts','Open-weight, custom license','Very long context (10M Scout/1M Maverick) and wide hosting ecosystem [S21]','Large weights; custom license; quality/host price differs','Yes','Host/model dependent','Host/model dependent','Not first capstone choice'),
('Hugging Face Providers/Endpoints','HF + providers','Hosting layer','One SDK to experiment across 200+ models; dedicated endpoints available [S22–S23]','Not one model; behavior, pricing and privacy depend on underlying provider','Yes','Varies','Varies','Experimentation/hosting option')]
table(d,['Model/API','Provider','Type','Strengths / verified facts','Weaknesses / limitations','RAG','Tools','Structured','HisabDo fit'],llm_rows,font=6.2)

heading(d,'3.1 Criterion-by-criterion interpretation',2)
table(d,['Criterion','Finding for HisabDo'],[
('Accuracy/reliability','No specification replaces a HisabDo test set. Use deterministic calculators and tools for numbers; score grounded answers, refusals and tool arguments.'),
('Financial reasoning','General reasoning is useful for explanations, but model arithmetic and personalized advice are not authoritative. Escalate complex educational questions; never write transactions without explicit application controls.'),
('Structured output','OpenAI, Gemini and current Claude models document schema/structured output support [S2,S7,S14]. Validate again with Pydantic; schema conformance is not semantic truth.'),
('Tool calling','All three major APIs support tools; open models may need serving-specific parsers/templates. Keep tools narrow, authenticated and read-only by default.'),
('RAG','Every candidate can generate from retrieved context. Retrieval quality, chunking, citation validation and prompt-injection defenses matter more than giant context windows.'),
('Latency','Small/Flash/Luna tiers should be tested under realistic prompts. Stream user-facing text with SSE, but do not stream unvalidated high-risk conclusions.'),
('Scalability','Hosted APIs remove GPU operations; local embeddings/classifier keep predictable cost. Add queues for ingestion and bounded concurrency for providers.'),
('Privacy','OpenAI API data is not used for training by default and standard abuse monitoring can retain content up to 30 days; eligibility varies by endpoint [S3]. Google’s unpaid services may use content; paid services have different terms [S10]. Anthropic commercial/API data is not used for training without permission; ZDR requires eligibility/arrangement [S15].'),
('Vendor lock-in','Use an internal ChatProvider interface and a provider-neutral response schema. Avoid provider-managed vector stores in the first version.'),
('Capstone fit','A hosted low-cost LLM + local retrieval + local classifier demonstrates architecture, evaluation and safety without operating a large GPU fleet.')
],font=7.5)

heading(d,'3.2 Weighted decision matrix',2)
para(d,'Scores (1 poor–5 strong) are recommendations, not verified provider benchmarks. Weighting favors capstone practicality. Re-score after the HisabDo evaluation set; a one-point change can alter rank.')
weights={'Quality':20,'Structured/tools':15,'RAG':10,'Cost':15,'Latency':10,'Privacy/control':10,'Ease':10,'Lock-in/fallback':10}
scores={
'GPT-5.6 Luna':[4,5,5,5,4,3,5,3],
'Gemini 3.1 Flash-Lite':[4,5,5,5,5,3,5,3],
'Gemini 3.7 Flash':[5,5,5,3,4,3,4,3],
'Claude Sonnet 5':[5,5,5,2,4,4,4,3],
'Mistral Small 4 API':[4,4,4,5,5,4,4,4],
'Qwen3 8B self-hosted':[3,3,4,4,3,5,2,5]
}
rows=[]
for m,vals in scores.items():
    total=sum(v*w for v,w in zip(vals,weights.values()))/5
    rows.append([m]+vals+[f'{total:.0f}/100'])
table(d,['Candidate']+list(weights.keys())+['Weighted'],rows,font=6.6)
para(d,'Interpretation: GPT-5.6 Luna is selected because its documented API contract, cost and integration are attractive and predictable. Gemini 3.1 Flash-Lite and Mistral Small 4 remain credible challengers; the tiny difference in paper scores is not statistically meaningful. The final choice must be based on HisabDo test results and account/region availability.')

heading(d,'4. RAG and Embedding Comparison',1)
rag_rows=[
('BAAI/bge-m3','Local open model','1024','8,192','100+ languages; dense, sparse and multi-vector; suited to hybrid retrieval [S26]','Local compute; larger/slower than MiniLM/E5-small','Primary'),
('multilingual-e5-small','Local open model','384','Verify model card/runtime','100-language base; compact 12-layer model [S27]','Lower-resource languages may degrade; query/passage prefixes required by E5 usage','Lightweight fallback'),
('all-MiniLM-L6-v2','Local open model','384','Short text','Very small (22.7M noted in efficiency docs), fast CPU/ONNX [S24–S25]','English-oriented; weaker fit for Urdu/Roman Urdu and long policy chunks','English-only prototype'),
('OpenAI text-embedding-3-small','Hosted API','Provider default/configurable','Verify current model page','$0.02/M tokens; easy API [S4]','Data leaves service boundary; vendor dependency','Low-cost hosted alternative'),
('OpenAI text-embedding-3-large','Hosted API','Provider default/configurable','Verify current model page','Most capable OpenAI embedding tier; $0.13/M [S5]','Higher cost than small/local','Quality challenger'),
('Gemini Embedding 2','Google hosted','128–3072; 768 recommended option','8,192','Stable multimodal; 100+ languages; auto-normalized truncated dimensions [S9]','Paid $0.20/M text on current page; API/data dependency [S6]','Multimodal alternative'),
('Voyage 4 / finance-2','Voyage hosted','1024 default/flexible on v4','32,000','Strong retrieval focus; finance-2 domain model; $0.12/M listed [S28–S29]','Extra vendor; provider says general v4 now outperforms domain models—must benchmark','Premium experiment')]
table(d,['Embedding','Deployment','Dimension','Input','Strengths / facts','Weaknesses','Decision'],rag_rows,font=6.7)
heading(d,'4.1 Recommended RAG pipeline',2)
numbered(d,[
'Ingest only approved HisabDo documents. Store document_id, title, version, effective_date, sensitivity, language and source URL/path.',
'Normalize Unicode but preserve merchant names, headings, numbers and tables. Detect English/Urdu/Roman-Urdu where possible.',
'Chunk by heading/section (target roughly 350–700 tokens with 10–15% overlap); never split policy exceptions from their conditions.',
'Embed documents with BGE-M3 locally; start with dense vectors. Add BM25/sparse retrieval when query-set results justify it.',
'Store in Qdrant local Docker. Qdrant is model-agnostic and has a managed free cluster option; current official page describes 0.5 vCPU, 1 GB RAM and 4 GB disk with inactivity rules [S30].',
'At query time, apply tenant/sensitivity/version filters, retrieve top 15–20, optionally rerank, then pass the best 4–8 chunks within a strict token budget.',
'Require citation objects containing document_id and chunk_id. Reject citations not in the retrieved set.',
'If retrieval confidence/evidence is insufficient, say that the knowledge base does not contain the answer and route to support.'
])
heading(d,'4.2 Embedding selection gate',2)
para(d,'Build at least 100 query–relevant-chunk pairs, including abbreviations, typos, Urdu script, Roman Urdu, category terms and adversarial distractors. Compare BGE-M3, multilingual-e5-small and one hosted model. Recommended gate: Recall@5 ≥ 0.90 on critical policy/app-help queries, no language slice more than 10 percentage points below overall, and citation precision ≥ 0.95 after generation. These thresholds are project decisions, not industry guarantees.')

heading(d,'5. Expense Categorization Analysis',1)
table(d,['Approach','Strengths','Weaknesses','When to use','HisabDo decision'],[
('Traditional ML: TF-IDF + Logistic Regression/Linear SVM','Fast CPU inference; cheap; interpretable features/errors; easy retraining; strong for recurring merchant text.','Needs labeled examples; concept drift; sparse signals for unseen merchants.','Default for controlled categories and short transaction descriptions.','Primary model.'),
('Transformer classifier','Better semantic generalization; can fine-tune multilingual encoder.','More data/compute/ops; calibration and latency work; harder to explain.','Upgrade when labeled set is large and baseline errors plateau.','Phase 2 challenger.'),
('LLM classification','Zero/few-shot start; handles novel text; explanations.','Variable, expensive, privacy exposure, nondeterministic, label drift and hallucination.','Bootstrapping labels or anonymized low-confidence review.','Not primary.'),
('Hybrid','Rules + ML + confidence gate + optional LLM/human review balances precision and coverage.','More routing logic and evaluation.','Realistic systems with known merchants and unknown long tail.','Recommended system design.')
],font=7.4)
heading(d,'5.1 Feature and model design',2)
bullets(d,[
'Inputs: normalized description, merchant, channel, currency, debit/credit flag, recurring indicator and optional amount bucket. Exclude user identity and exact account number.',
'Rules first: user overrides, known merchant map, transfers, salary/refund patterns and internal transactions. Version the rule set.',
'ML features: word ngrams (1–2), character ngrams (3–5), normalized merchant token, channel and amount bucket. Character ngrams handle spelling variation and abbreviated merchant strings.',
'Model: multinomial Logistic Regression with class weights when imbalance is material. Compare calibrated LinearSVC and a simple most-frequent/merchant baseline.',
'Output: category_id from a fixed taxonomy, confidence, source (rule/model/fallback), top alternatives, model_version and needs_review.',
'Learning loop: preserve user corrections as labeled events, de-duplicate correlated transactions, review labels, and retrain on a time-based schedule.'
])
heading(d,'5.2 Evaluation and release gates',2)
para(d,'Use a time-based split where possible so repeat transactions from the same user/merchant do not leak across train and test. Report accuracy, macro-F1, weighted-F1, per-class precision/recall, confusion matrix and coverage at confidence thresholds. Scikit-learn documents text pipelines and classification reports; F1 supports multiclass averaging [S35–S36].')
table(d,['Metric/gate','Proposed capstone target','Reason'],[
('Macro-F1','≥ 0.80 and ≥ baseline + 0.10','Protects small classes; target must be adjusted to taxonomy/data quality.'),
('High-confidence precision','≥ 0.95 at selected threshold','Automatic labels should be trustworthy.'),
('Coverage','Report at 0.95 precision; do not hide abstentions','Shows practical usefulness of confidence gating.'),
('Critical confusion','Transfers/income/refunds must have explicit review','These errors can distort financial summaries.'),
('Latency','p95 < 50 ms in-process on target CPU, excluding network','Traditional model should be effectively instant.'),
('Drift','Alert on unknown-token rate, confidence shift and category distribution','Merchant behavior changes over time.')
],font=7.8)

heading(d,'6. FastAPI Architecture Recommendation',1)
para(d,'Logical flow:')
callout(d,'Architecture','HisabDo Frontend → FastAPI API boundary → AI Orchestrator → [Chat Provider Adapter | Expense Classifier | RAG Retriever] → Qdrant / HisabDo read-only tools → Output & Policy Validation → Frontend')
heading(d,'6.1 Components',2)
table(d,['Component','Responsibility'],[
('API routers','/v1/assistant/chat, /v1/assistant/chat/stream, /v1/expenses/categorize, /v1/feedback, /health/ready. Authenticate and rate-limit.'),
('Pydantic contracts','Strict enums/lengths/ranges; forbid extra fields for sensitive contracts; generate OpenAPI; validate provider output again [S34].'),
('Assistant orchestrator','Risk/intent classification, retrieval, approved tool calls, prompt assembly, retries, fallback and redacted telemetry.'),
('Provider adapters','One internal interface for OpenAI/Gemini/Mistral/local OpenAI-compatible endpoint. Normalize errors, usage and citations.'),
('RAG service','Chunk lookup with tenant and version filters; hybrid search; optional rerank; citation allow-list.'),
('Expense service','Load immutable model artifact at startup; rules → classifier → threshold; no network call on normal path.'),
('Validators/guardrails','Schema, citation, numeric and policy validation; prompt-injection checks; disclaimer/escalation rules.'),
('Observability','request_id, model/version, latency, tokens, cost estimate, retrieval IDs, route/fallback and validation result; redact content.'),
('Background workers','Document ingestion, embedding, evaluation and model retraining; keep these out of request latency path.')
],font=7.6)
heading(d,'6.2 Reliability pattern',2)
bullets(d,[
'Use asynchronous HTTP clients for hosted APIs, explicit connect/read/total timeouts, bounded retries with jitter only for retryable failures, and a circuit breaker.',
'Return 422 for invalid client input, 401/403 for auth, 429 for local throttling, 503 for exhausted upstream capacity, and a stable application error envelope. FastAPI supports custom validation and exception handlers [S31].',
'Stream safe low-risk assistant text using SSE [S32]. For high-risk responses, buffer until citation and policy validation complete.',
'Cache only non-personal FAQ/RAG answers by knowledge-base version and normalized query. Never cross-cache user-specific account data.',
'Use idempotency keys for any future write action. The recommended first version exposes read-only tools only.',
'Deploy with separate secrets, least-privilege database credentials, HTTPS, OAuth2/JWT scopes, and per-user/tenant filters [S33].'
])
heading(d,'6.3 Internal response contract',2)
code='''AssistantResponse\n- answer: str\n- intent: Literal[app_help, education, account_query, high_risk, unsupported]\n- citations: list[{document_id, chunk_id, title}]\n- tool_results_used: list[str]\n- confidence: float (0..1; operational signal, not probability of truth)\n- needs_human_review: bool\n- disclaimer: str | null\n- provider: str\n- model: str\n- request_id: UUID'''
p=d.add_paragraph(); p.style='Intense Quote'; p.add_run(code).font.name='Consolas'

heading(d,'7. Final Recommendations',1)
heading(d,'7.1 Primary recommendation',2)
callout(d,'Primary stack','GPT-5.6 Luna + BGE-M3 + Qdrant + rules/TF-IDF Logistic Regression + FastAPI/Pydantic.')
bullets(d,[
'Why: a strong current developer feature set and low published token price for the assistant [S1–S2], no third-party embedding cost, multilingual/hybrid retrieval potential [S26], and a classifier that is fast and objectively measurable.',
'Control: Luna handles language and tool selection; HisabDo tools supply authoritative data. Escalate to GPT-5.6 Terra only for approved, complex, non-transactional questions after cost controls.',
'Condition: must pass HisabDo quality/safety gates. If Gemini or Mistral wins the same frozen evaluation at materially lower cost/latency, switch through the adapter.'
])
heading(d,'7.2 Secondary recommendation',2)
callout(d,'Secondary stack','Paid Gemini 3.1 Flash-Lite (or evaluate Gemini 3.7 Flash) + the same local BGE-M3/Qdrant/classifier components.')
bullets(d,[
'Why: stable Flash-Lite has documented structured output, function calling, file/search capabilities and a 1M input context [S7]. It is an inexpensive independent vendor route.',
'Privacy condition: do not use unpaid Gemini services with real personal/financial data because current terms state unpaid content may be used to improve products and human reviewers may process it [S10]. Activate billing and review regional/ZDR requirements.',
'Use: integration test and optional failover after the primary implementation is stable; avoid paying the capstone complexity cost of active-active multi-provider routing on day one.'
])
heading(d,'7.3 Fallback recommendation',2)
callout(d,'Fallback','Deterministic FAQ/RAG extracts and calculators first; optional local Qwen3 8B only for non-authoritative wording.')
bullets(d,[
'If hosted LLMs fail, return retrieved excerpts with citations, application help links and deterministic calculator/tool outputs rather than a fabricated answer.',
'Qwen documentation supports vLLM/SGLang OpenAI-compatible serving and tool-call parsing [S19–S20], but local hardware and lower assurance make it a degraded, non-high-stakes fallback—not the source of ledger facts.',
'Expense categorization remains available because it is local and independent of the chatbot API.'
])

heading(d,'8. Cost and Practicality',1)
para(d,'Current verified examples: OpenAI lists Luna at $0.20/M input and $1.20/M output [S1]. Anthropic lists Sonnet 5 at $2/$10 and Haiku 4.5 at $1/$5 [S13]. Mistral’s current pricing page lists Small 4 at $0.15/$0.60 [S17]. Google prices vary by model and service class and include temporary rates; the current official pricing page must be consulted [S6,S8]. Hugging Face Inference Providers gives very small monthly experimentation credits and dedicated endpoints are charged by compute time [S22–S23].')
para(d,'Illustrative workload (not a quote): 1,000 assistant turns/month, each with 2,000 billed input tokens and 500 output tokens; no caching, reasoning surcharge, search, file storage, taxes or retries. Formula: turns × ((input_tokens/1M × input_rate) + (output_tokens/1M × output_rate)).')
price_rows=[]
for name,ip,op in [('GPT-5.6 Luna',.20,1.20),('GPT-5.6 Terra',2,12),('Claude Sonnet 5',2,10),('Claude Haiku 4.5',1,5),('Mistral Small 4',.15,.60),('Gemini 3.7 Flash promo',.75,3.75)]:
    c=1000*((2000/1e6)*ip+(500/1e6)*op)
    price_rows.append((name,f'${ip:.2f}',f'${op:.2f}',f'${c:.2f}'))
table(d,['Model','Input / M','Output / M','Illustrative monthly token cost'],price_rows,font=8)
bullets(d,[
'Free tiers are suitable only for synthetic/demo data and availability is not a service guarantee. Google’s current pricing table explicitly distinguishes whether data is used to improve products [S6].',
'Local BGE-M3 and the expense classifier have no per-token fee but consume CPU/RAM and engineering time. Qdrant local is free software; managed cloud pricing and inactivity limits must be checked [S30].',
'Self-hosting a capable LLM replaces token charges with GPU rental, operations, patching, monitoring and lower utilization. It is usually less practical for a modest capstone chatbot than a hosted low-cost API.',
'Add per-user quotas, maximum prompt/retrieval/output sizes, cost telemetry, monthly budget alerts, response caching for public FAQ, and an escalation budget.'
])

heading(d,'9. Risks and Limitations',1)
table(d,['Risk','Impact','Mitigation'],[
('Hallucination / financial misinformation','Wrong explanations may influence user decisions.','RAG citations; deterministic tools; abstain; disclaimers; high-risk intent policy; golden-set evaluation; user report action.'),
('Prompt injection','Retrieved or user text attempts to call tools/exfiltrate data.','Treat content as data, not instructions; allow-listed tools; argument validation; least privilege; citation-only context; adversarial tests.'),
('Privacy','Financial records leave the trust boundary or enter logs.','Minimize/redact; paid commercial terms; no free-tier real data; encryption; retention policy; ZDR review; content-free telemetry.'),
('API dependency / availability','Outage or rate limit blocks chat.','Timeouts, circuit breaker, deterministic fallback, provider adapter, status monitoring.'),
('Cost growth','Long histories/RAG and retries multiply tokens.','Budgets, truncation/summarization, cache, quotas, token accounting, route routine tasks to local models/classifier.'),
('Latency','RAG + model + tool calls delay UI.','Parallel safe calls, precomputed embeddings, bounded top-k, streaming after validation, p95 SLO.'),
('Model change/deprecation','Alias behavior changes or model retires.','Pin versions where supported, regression suite, deprecation calendar, canary rollout, adapter.'),
('Vendor lock-in','Provider-specific prompts/tools complicate migration.','Internal schemas, own Qdrant index, provider-neutral service interface, contract tests.'),
('Security','Key leakage, broken access control, injection, excessive tool permission.','Secret manager, scopes, tenant filters, audit logs, dependency scans, egress allow-list, read-only tools.'),
('Classifier bias/drift','Rare merchants/languages miscategorized.','Macro-F1 by language/category, confidence gate, correction loop, periodic drift review.'),
('Evaluation limitation','Vendor benchmarks do not equal HisabDo quality.','Frozen representative data, blind human review, paired tests, error taxonomy and release gates.')
],font=7.1)

heading(d,'10. Final Technology Stack Recommendation',1)
table(d,['Layer','Recommended technology','Alternative / trigger'],[
('LLM','OpenAI GPT-5.6 Luna via Responses API; low reasoning for routine use; Terra only by policy.','Paid Gemini 3.1 Flash-Lite; benchmark Gemini 3.7 Flash and Mistral Small 4.'),
('Embedding','BAAI/bge-m3, local, 1024-d dense first.','multilingual-e5-small for lower RAM/latency; Gemini Embedding 2 for multimodal hosted retrieval.'),
('Vector DB','Qdrant local Docker; payload filters and collection aliases.','Qdrant Cloud when deployment/backup/availability justify it; PostgreSQL pgvector if already standardized.'),
('Expense classifier','Rules + TF-IDF word/character ngrams + Logistic Regression; confidence review.','Fine-tuned multilingual transformer only after data/benchmark evidence.'),
('Backend','FastAPI, async provider clients, versioned routers.','None required.'),
('Validation','Pydantic v2 request, provider-output and response schemas; domain validators.','JSON Schema generated from Pydantic for provider structured output.'),
('Evaluation','pytest + scikit-learn; Recall@k/MRR/citation precision; custom groundedness and safety set; optional Ragas.','Any framework is acceptable if metrics and datasets remain portable.'),
('Observability','OpenTelemetry-compatible traces/metrics, structured redacted logs, token/cost ledger.','Capstone may start with JSON logs + Prometheus metrics.'),
('Fallback','Retrieved excerpts + deterministic calculators/FAQ; optional local Qwen3 8B for wording.','Second hosted provider after contract tests.'),
('Deployment','One FastAPI container + Qdrant + background worker; classifier and embedder loaded at startup.','Split AI services only after measured scaling need.')
],font=7.4)

heading(d,'Recommendation for HisabDo Team',1)
callout(d,'Exact technologies','GPT-5.6 Luna; BAAI/bge-m3; Qdrant; rules + scikit-learn TF-IDF Logistic Regression; FastAPI; Pydantic v2; pytest/scikit-learn retrieval and safety evaluation; paid Gemini 3.1 Flash-Lite as the secondary provider; deterministic extract/tool fallback plus optional local Qwen3 8B.')
para(d,'Why selected: this split gives the assistant a modern structured/tool-capable API while keeping retrieval and transaction classification local, inexpensive and testable. It avoids operating a large LLM GPU, limits vendor lock-in, supports multilingual experiments, and provides useful degraded behavior during an API outage.')
para(d,'Alternatives: choose Gemini 3.1 Flash-Lite if it passes the frozen HisabDo evaluation with equal safety and lower measured cost/latency; choose Mistral Small 4 for a low-cost/open-weight path; choose multilingual-e5-small if BGE-M3 is too slow; choose a fine-tuned transformer classifier only after the baseline plateaus on sufficient labeled data.')
heading(d,'Next implementation steps (ordered)',2)
numbered(d,[
'Freeze the expense taxonomy, high-risk assistant policy, allowed tools and Pydantic API contracts.',
'Create three versioned evaluation assets: 150 assistant questions, 100+ RAG query–chunk judgments, and a leakage-safe labeled transaction set.',
'Implement the rule + TF-IDF Logistic Regression baseline; publish macro-F1, per-class metrics, confidence/coverage curve and confusion matrix.',
'Ingest a small approved knowledge base with BGE-M3 into local Qdrant; measure Recall@5 and citation precision, including Urdu/Roman-Urdu slices.',
'Implement ChatProvider and EmbeddingProvider interfaces; connect GPT-5.6 Luna with structured output and read-only tools.',
'Add output/citation/numeric validators, timeouts, retries, circuit breaker, quotas and safe deterministic fallback.',
'Run a blinded GPT-5.6 Luna vs paid Gemini 3.1 Flash-Lite (and optional Mistral Small 4) evaluation. Select on quality gates first, then cost/latency.',
'Perform prompt-injection, tenant-isolation, key-handling and log-redaction tests before using any real financial data.',
'Deploy a monitored capstone pilot with synthetic or consented data; record model IDs, prices, terms and evaluation results in an architecture decision record.'
])

heading(d,'Appendix A — Acceptance Gates',1)
table(d,['Area','Block release if…'],[
('Assistant','Any test causes an unauthorized tool call, fabricated account fact, uncited KB claim in citation-required mode, or unsafe personalized high-risk recommendation.'),
('RAG','Critical Recall@5 is below the agreed threshold, citation IDs are not validated, or tenant/version filters can be bypassed.'),
('Categorization','Critical transfer/income/refund errors are not review-gated, data leakage exists, or model artifact/taxonomy versions are absent.'),
('API','Secrets appear in client bundles/logs, upstream calls lack timeouts, validation can be bypassed, or errors expose internal detail.'),
('Privacy','Real data is sent through an unpaid/consumer service without approved terms, consent and retention review.'),
('Operations','No cost cap, rate limit, fallback, model version inventory or rollback path exists.')
],font=7.6)
heading(d,'Appendix B — Research Limitations',1)
bullets(d,[
'Provider documentation is authoritative for the provider’s own current specification, not comparative financial accuracy.',
'Model availability, aliases, previews, rates, quotas and terms can change after the cut-off. This report is a snapshot dated '+DATE+'.',
'No HisabDo dataset, taxonomy, traffic profile, cloud budget or compliance requirement was supplied. Cost examples and thresholds therefore use explicit assumptions.',
'Pakistan-specific regulatory, consumer-protection, tax and data-residency requirements require qualified legal/compliance review; this research is technical, not legal advice.'
])
refs(d)
main_path=OUT/'01_HisabDo_AI_Models_APIs_Technical_Research.docx'; d.save(main_path)

# ---------------- Decision matrix / evidence workbook in Word ----------------
d=base_doc('Decision Matrix, Cost Scenarios & Evidence Register','Companion document for review, scoring updates and procurement verification')
add_toc_note(d)
heading(d,'1. How to Use This Companion',1)
para(d,'This document isolates the decision logic from the narrative report. Update the yellow decision inputs after a proof of concept; do not change verified facts without re-checking the cited official page. Scores are ordinal engineering estimates, not model benchmarks.')
heading(d,'2. Requirements and Weights',1)
req=[('Grounded answer/tool correctness',20,'Financial facts must come from RAG/tools.'),('Structured output and tools',15,'Reliable FastAPI contracts and narrow tools.'),('Cost',15,'Capstone budget and scaling.'),('Latency',10,'Interactive experience.'),('Privacy/control',10,'Financial data sensitivity.'),('Ease/maintainability',10,'Small student team.'),('RAG integration',10,'Knowledge-base answers and citations.'),('Lock-in/fallback',10,'Model churn and outages.')]
table(d,['Criterion','Weight %','Rationale'],req,font=8)
heading(d,'3. Candidate Scorecard',1)
rows=[]
for m,vals in scores.items():
    total=sum(v*w for v,w in zip(vals,weights.values()))/5
    rows.append([m]+vals+[f'{total:.0f}'])
table(d,['Candidate']+list(weights.keys())+['/100'],rows,font=6.8)
para(d,'Decision note: select only among candidates that pass hard safety/quality gates. Weighted score breaks ties; it must not compensate for a fabricated ledger fact or unsafe tool call.')
heading(d,'4. Verified Model/API Facts at Cut-off',1)
table(d,['Candidate','Availability/specification snapshot','Published price snapshot','Evidence'],[
('GPT-5.6 Luna','1.05M context; 128k max output; functions, web/file search.','USD 0.20/M input, 1.20/M output.','[S1]'),
('GPT-5.6 Terra','1.05M context; 128k max output; same listed tools.','USD 2/M input, 12/M output.','[S1]'),
('Gemini 3.1 Flash-Lite','Stable; 1,048,576 input and 65,536 output; functions/structured output/file/search.','Pricing has service classes; verify current [S6].','[S6,S7]'),
('Gemini 3.7 Flash','1M context, 64k output, tunable thinking.','Promo standard USD 0.75/M input, 3.75/M output through 31 Dec 2026; standard rises after.','[S8]'),
('Claude Sonnet 5','1M context; 128k output; fast balanced tier.','USD 2/M input, 10/M output.','[S12,S13,S14]'),
('Claude Haiku 4.5','200k context; 64k output.','USD 1/M input, 5/M output.','[S12,S13]'),
('Mistral Small 4','Current guide lists functions and structured outputs.','USD 0.15/M input, 0.60/M output on current pricing page.','[S16,S17]'),
('Qwen3 local','vLLM/SGLang OpenAI-compatible serving; Hermes-style tool parsing.','No token API price when self-hosted; infrastructure is not free.','[S19,S20]')
],font=7)
heading(d,'5. Cost Calculator',1)
para(d,'Inputs: N turns, I average billed input tokens/turn, O output tokens/turn. Monthly token cost = N × ((I/1,000,000 × input rate) + (O/1,000,000 × output rate)). This excludes caching, thinking/reasoning treatment, tools, search, retries, storage, hosting, network, taxes and currency conversion.')
scenarios=[('Demo',100,1000,300),('Capstone pilot',1000,2000,500),('Small production',10000,2500,600)]
models=[('GPT-5.6 Luna',.2,1.2),('GPT-5.6 Terra',2,12),('Claude Sonnet 5',2,10),('Claude Haiku 4.5',1,5),('Mistral Small 4',.15,.6),('Gemini 3.7 Flash promo',.75,3.75)]
rows=[]
for s,n,i,o in scenarios:
    for m,ip,op in models:
        c=n*((i/1e6)*ip+(o/1e6)*op)
        rows.append((s,m,f'{n:,}',f'{i:,}',f'{o:,}',f'${c:,.2f}'))
table(d,['Scenario','Model','Turns','Input/turn','Output/turn','Token cost'],rows,font=7)
heading(d,'6. Embedding Decision Matrix',1)
table(d,['Candidate','Quality potential','Multilingual','Local/privacy','Speed/resource','Lock-in','Recommended role'],[
('BGE-M3','High retrieval flexibility','100+ languages','Yes','Moderate/heavier','Low','Default; benchmark first.'),
('multilingual-e5-small','Good compact baseline','100-language base','Yes','Fast/light','Low','Resource-constrained alternative.'),
('all-MiniLM-L6-v2','Good English compact baseline','Limited relative fit','Yes','Very fast/light','Low','English-only prototype.'),
('OpenAI embedding-3-small','Hosted baseline','Provider states non-English capable family','No','Fast API','Medium','Cheap hosted challenger.'),
('Gemini Embedding 2','Multimodal and flexible dimensions','100+ languages','No','Fast API','Medium','Multimodal future option.'),
('Voyage 4 / finance-2','Retrieval-specialized','Model dependent','No','Fast API','Medium','Optional premium experiment.')
],font=7.2)
heading(d,'7. Expense Model Experiment Register',1)
table(d,['Experiment','Features/model','Required report','Advance if…'],[
('B0','Most-frequent category','Accuracy, macro-F1','Reference only.'),
('B1','Merchant exact/regex rules','Coverage, precision, conflicts','Very high precision on covered items.'),
('M1','Word TF-IDF + Logistic Regression','Macro/per-class F1, confusion, calibration','Beats B0 materially.'),
('M2','Word + char TF-IDF + Logistic Regression','Same + latency + selective curve','Recommended if better macro-F1/robustness.'),
('M3','Calibrated LinearSVC','Same + calibration error','Adopt only if selective accuracy improves.'),
('T1','Fine-tuned multilingual transformer','Same + RAM/latency/ops','Advance only if gain justifies complexity.'),
('L1','LLM zero/few-shot, anonymized sample','Accuracy, consistency, cost, privacy review','Use only as low-confidence/human-review aid.')
],font=7.3)
heading(d,'8. Evaluation Worksheet',1)
table(d,['Suite','Minimum cases','Metrics','Hard failures'],[
('Assistant','150, including normal, ambiguous, adversarial and high-risk','Grounded correctness, tool accuracy, schema pass, refusal, p95, cost','Fabricated ledger fact; unauthorized tool; unsafe advice.'),
('RAG','100+ judged queries and distractors','Recall@1/5, MRR, nDCG, citation precision, no-answer accuracy','Cross-tenant result; invented citation; critical miss.'),
('Expense','Time-split labeled set; enough support per class','Macro/weighted F1, per-class PR, selective accuracy, coverage, p95','Leakage; critical category auto-label error.'),
('Resilience','Timeout, 429, 5xx, malformed JSON, vector outage','Fallback correctness, error contract, recovery time','Crash, secret leak, unbounded retry.'),
('Security','Injection, authz, tenant, oversized input, logs','Pass/fail + evidence','Unauthorized data/tool access or sensitive log content.')
],font=7.1)
heading(d,'9. Architecture Decision Record (ADR) Template',1)
for label,text in [
('Decision','Use GPT-5.6 Luna as initial assistant default; BGE-M3/Qdrant for RAG; rules + TF-IDF Logistic Regression for expenses.'),
('Status','Proposed until proof-of-concept gates pass.'),
('Context','Capstone constraints: small team, low budget, need for measurable reliability and straightforward FastAPI integration.'),
('Consequences','Hosted LLM dependency remains; embeddings/classification stay local; a provider adapter and evaluation suite are mandatory.'),
('Revisit triggers','Provider deprecation/price change; quality gate failure; p95 SLO breach; privacy requirement change; >3 months classifier drift; knowledge base becomes multimodal.')]:
    para(d,f'{label}: {text}',bold_prefix=label+':')
refs(d)
matrix_path=OUT/'02_HisabDo_Decision_Matrix_Cost_Evidence.docx'; d.save(matrix_path)

# ---------------- FastAPI blueprint ----------------
d=base_doc('FastAPI AI Integration Blueprint','Implementation-ready service boundaries, contracts, controls and test plan')
add_toc_note(d)
heading(d,'1. Target Architecture',1)
callout(d,'Request path','Frontend → /v1 FastAPI router → authentication/rate limit → orchestrator → RAG/tool/classifier/provider adapter → semantic validation → typed response → redacted telemetry.')
heading(d,'1.1 Suggested project structure',2)
structure='''app/\n  main.py\n  api/v1/{assistant.py, expenses.py, feedback.py, health.py}\n  core/{config.py, security.py, errors.py, telemetry.py}\n  schemas/{assistant.py, expense.py, errors.py}\n  services/{assistant.py, rag.py, expense.py, validation.py}\n  providers/{base.py, openai_provider.py, gemini_provider.py, local_provider.py}\n  tools/{registry.py, balances.py, budgets.py, transactions.py}\n  repositories/{qdrant.py, knowledge.py}\n  ml/{pipeline.joblib, taxonomy.json, metadata.json}\n  tests/{contracts, unit, integration, safety, evals}/'''
p=d.add_paragraph(); p.style='Intense Quote'; r=p.add_run(structure); r.font.name='Consolas'; r.font.size=Pt(8.5)
heading(d,'2. API Contracts',1)
table(d,['Endpoint','Request','Response','Notes'],[
('POST /v1/assistant/chat','message ≤ 4,000 chars; conversation_id?; locale; idempotency/request id','AssistantResponse','Authenticated; no provider fields exposed unless useful for audit.'),
('POST /v1/assistant/chat/stream','Same','SSE events: meta, token, citation, final, error','Stream only after risk policy; final event contains validated object [S32].'),
('POST /v1/expenses/categorize','description, merchant?, amount?, currency, channel?, debit_credit','ExpensePrediction','Batch endpoint may be added for imports.'),
('POST /v1/feedback','request_id, rating, correction?, reason?','Accepted','No full transcript required by default.'),
('GET /health/live','None','status','Process health only.'),
('GET /health/ready','None','component readiness','Check loaded classifier and required dependencies without expensive provider generation.')
],font=7.5)
heading(d,'2.1 Pydantic model sketch',2)
code='''from enum import Enum\nfrom pydantic import BaseModel, ConfigDict, Field\n\nclass Intent(str, Enum):\n    app_help = "app_help"\n    education = "education"\n    account_query = "account_query"\n    high_risk = "high_risk"\n    unsupported = "unsupported"\n\nclass Citation(BaseModel):\n    model_config = ConfigDict(extra="forbid")\n    document_id: str = Field(min_length=1, max_length=128)\n    chunk_id: str = Field(min_length=1, max_length=128)\n    title: str = Field(max_length=300)\n\nclass AssistantResponse(BaseModel):\n    model_config = ConfigDict(extra="forbid")\n    answer: str = Field(min_length=1, max_length=8000)\n    intent: Intent\n    citations: list[Citation] = Field(max_length=12)\n    tool_results_used: list[str] = Field(max_length=10)\n    confidence: float = Field(ge=0, le=1)\n    needs_human_review: bool\n    disclaimer: str | None = Field(default=None, max_length=1000)\n    request_id: str'''
p=d.add_paragraph(); p.style='Intense Quote'; r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(7.4)
para(d,'Generate provider structured-output JSON Schema from the Pydantic model, then validate the returned object again. Pydantic can validate JSON and generate JSON Schema [S34]; provider schema compliance does not validate citations or financial meaning.')
heading(d,'3. Provider Adapter',1)
code='''class ChatProvider(Protocol):\n    async def generate(\n        self, *, messages, tools, output_schema, timeout_s, request_id\n    ) -> ProviderResult: ...\n\nclass ProviderResult(BaseModel):\n    parsed: dict\n    model: str\n    input_tokens: int | None\n    output_tokens: int | None\n    latency_ms: int\n    finish_reason: str | None\n    raw_request_id: str | None'''
p=d.add_paragraph(); p.style='Intense Quote'; r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(8)
bullets(d,[
'Keep provider prompts and parsing inside each adapter; expose only internal messages, tools and schemas.',
'Normalize provider errors into timeout, rate_limited, unavailable, invalid_output, safety_refusal and authentication_error.',
'Configuration: PRIMARY_CHAT_PROVIDER, PRIMARY_CHAT_MODEL, SECONDARY_CHAT_PROVIDER, MAX_OUTPUT_TOKENS, TIMEOUT_SECONDS and monthly budget.',
'Use model aliases only in configuration. Record the actual returned model/version when the provider exposes it.',
'Do not automatically fail over after a provider may have executed a write tool. Recommended v1 tools are read-only.'
])
heading(d,'4. Assistant Orchestration',1)
numbered(d,[
'Validate input and authenticate user/tenant.',
'Redact or tokenize unnecessary identifiers; assign request_id.',
'Classify risk/intent with deterministic patterns plus small model/LLM only when needed.',
'For knowledge questions, retrieve from Qdrant with tenant, language, document status and version filters.',
'For account questions, expose narrowly scoped read-only functions such as get_budget_summary(start,end), list_recent_transactions(limit,filters) and calculate_spend_by_category(period).',
'Call the provider with a strict system policy, retrieved evidence, tool schema and AssistantResponse JSON Schema.',
'Validate tool arguments before execution; cap rows/date range and authorize every call independently.',
'Validate parsed output, citations, number provenance, disclaimer and risk policy. Retry repair once only if safe.',
'On failure, return retrieved excerpts/tool result in a deterministic template and a stable degraded flag.',
'Emit content-minimized telemetry and cost estimate; never log API keys, authorization headers or full account data.'
])
heading(d,'5. RAG Implementation',1)
table(d,['Stage','Implementation choice','Validation'],[
('Ingestion','Background job; approved document registry; heading-aware chunks; immutable document version.','Checksum, source, effective date, language, sensitivity.'),
('Embedding','BGE-M3 local; batch encoding; fixed normalized vector configuration.','Model/hash/dimension stored in collection metadata.'),
('Index','Qdrant collection alias knowledge_active; payload indexes for tenant/status/language.','Dimension/distance check; no cross-tenant query.'),
('Retrieve','Dense top 20, optional sparse/BM25 fusion, rerank to 6.','Frozen relevance set; Recall@5/MRR.'),
('Prompt','Quote chunk boundaries and IDs; state that documents are untrusted data.','Token budget; remove unsupported metadata.'),
('Output','Citations must be subset of retrieved chunk IDs.','Citation precision and claim support review.'),
('No-answer','Threshold + model abstention.','Measure false-answer rate separately.')
],font=7.2)
heading(d,'6. Expense Service',1)
code='''def categorize(tx: ExpenseRequest) -> ExpensePrediction:\n    # 1. user override / merchant rules\n    if rule_hit := rules.match(tx):\n        return prediction(rule_hit.category, 1.0, source="rule")\n    # 2. local pipeline: preprocessing + TF-IDF + LogisticRegression\n    probabilities = pipeline.predict_proba([feature_text(tx)])[0]\n    top = top_k(probabilities, k=3)\n    # 3. selective prediction\n    return ExpensePrediction(\n        category_id=top[0].label if top[0].score >= threshold else "needs_review",\n        confidence=top[0].score, alternatives=top,\n        source="model", model_version=MODEL_VERSION,\n        needs_review=top[0].score < threshold,\n    )'''
p=d.add_paragraph(); p.style='Intense Quote'; r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(7.8)
bullets(d,[
'Load joblib artifact once during application lifespan; never unpickle an untrusted artifact.',
'Store taxonomy version, training data window/hash, package versions, threshold and metrics beside the artifact.',
'Return 503 readiness failure if the artifact cannot load; do not silently switch to random/LLM classification.',
'Corrections are events, not direct model mutations. Review and retrain offline; promote through a model registry or signed artifact process.'
])
heading(d,'7. Error and Fallback Matrix',1)
table(d,['Failure','API behavior','User-safe response','Operational action'],[
('Invalid input','422 structured validation error','Ask to correct the specific field.','Count only; no upstream call [S31].'),
('Provider timeout/5xx','Retry once if safe, then degraded 200 or 503 by endpoint contract','Show retrieved/tool-backed answer or service-unavailable message.','Circuit breaker, alert on rate.'),
('Provider 429','No retry storm; respect retry hint','Try later; optional configured secondary.','Quota/budget alert.'),
('Malformed/schema-invalid output','One bounded repair; then fallback','Do not show raw model output.','Capture redacted failure category.'),
('RAG unavailable','Skip KB claims','Account tools/FAQ link only; disclose temporary limitation.','Vector readiness alert.'),
('Low retrieval evidence','No generation of KB claim','“I could not find this in the approved knowledge base.”','Add query to evaluation backlog.'),
('Expense confidence low','Return needs_review','Ask user to select among top categories.','Store correction after consent.'),
('Classifier missing','503 categorize endpoint','Categorization temporarily unavailable.','Fail readiness; rollback artifact.')
],font=7.1)
heading(d,'8. Security and Privacy Checklist',1)
bullets(d,[
'Keys only on server; secret manager/environment injection; rotate and scope; egress allow-list to selected providers.',
'JWT/OAuth2 scopes and independent authorization at every tool/repository call; FastAPI documents scope validation [S33].',
'Per-tenant filters are mandatory server-side parameters, never model-supplied values.',
'Prompt injection tests for user text and retrieved documents; tool names/arguments allow-listed; no arbitrary SQL or URLs.',
'Maximum body, message, history, retrieval, tool-result and output sizes; MIME allow-list and malware scan for future uploads.',
'Use paid/commercial API terms for personal data; document provider region, retention and ZDR eligibility [S3,S10,S11,S15].',
'Log request metadata, not raw financial content. Hash stable identifiers with a rotated secret if correlation is needed.',
'Encrypt in transit and at rest; define deletion, backup and incident-response procedures; perform dependency/SBOM scans.'
])
heading(d,'9. Test Plan',1)
table(d,['Test class','Examples','Pass condition'],[
('Unit','Pydantic bounds, rules, normalization, confidence threshold, citation subset.','Deterministic; branch coverage on critical validators.'),
('Contract','Each provider adapter against recorded/sandbox responses; OpenAPI snapshot.','Same internal result/error types.'),
('Integration','FastAPI + Qdrant + mock provider + model artifact.','End-to-end IDs, filters, fallbacks and statuses.'),
('Safety','Prompt injection, high-risk advice, unauthorized tool, fabricated citation.','No dangerous action/claim; correct escalation.'),
('Privacy','Log inspection, headers, exception paths, feedback.','No secret/financial content outside approved storage.'),
('Load','Assistant concurrency, categorize batch, Qdrant search.','p95 SLO and bounded resources; no retry amplification.'),
('Evaluation','Frozen assistant/RAG/expense datasets.','Meets release gates; report archived with model/prompt/index versions.'),
('Chaos','Provider timeout/429/5xx, Qdrant down, bad artifact.','Stable fallback/error contract and readiness behavior.')
],font=7.1)
heading(d,'10. Delivery Plan',1)
table(d,['Sprint','Deliverable','Exit evidence'],[
('1','Contracts, taxonomy, policies, evaluation sets, local stack.','Reviewed schemas and frozen test assets.'),
('2','Expense baseline + feedback endpoint.','Metrics, confusion matrix, threshold/coverage report.'),
('3','BGE-M3/Qdrant ingestion and retrieval.','Recall@k/MRR and multilingual slice report.'),
('4','GPT-5.6 Luna adapter, read-only tools, validators, fallback.','Assistant safety/quality/latency/cost report.'),
('5','Gemini/Mistral challenger evaluation; hardening.','Paired comparison, security checklist, ADR.'),
('6','Pilot deployment and supervisor demo.','Runbook, architecture diagram, monitored demo and rollback.')
],font=7.4)
heading(d,'Recommendation for HisabDo Team',1)
para(d,'Implement one well-tested primary path first: GPT-5.6 Luna through a ChatProvider adapter, local BGE-M3/Qdrant retrieval, and an in-process TF-IDF Logistic Regression categorizer. Keep all account truth in deterministic, authorized tools. Add paid Gemini 3.1 Flash-Lite only after contract and evaluation tests pass; preserve a deterministic no-LLM fallback.')
refs(d,['S1','S2','S3','S6','S7','S10','S11','S15','S19','S20','S26','S30','S31','S32','S33','S34','S35','S36'])
blue_path=OUT/'03_HisabDo_FastAPI_AI_Integration_Blueprint.docx'; d.save(blue_path)

# README
readme=f'''HisabDo AI Models/APIs Research Package\nResearch cut-off: {DATE}\n\nFiles\n1. 01_HisabDo_AI_Models_APIs_Technical_Research.docx\n   Main submission document: executive summary, requirements, comparisons, RAG, expense categorization, architecture, recommendations, costs, risks and final stack.\n\n2. 02_HisabDo_Decision_Matrix_Cost_Evidence.docx\n   Companion scoring matrix, pricing scenarios, experiment register, evaluation worksheet and ADR template.\n\n3. 03_HisabDo_FastAPI_AI_Integration_Blueprint.docx\n   Implementation-facing routes, contracts, provider adapter, orchestration, fallback, security and test plan.\n\nImportant\n- Prices/model availability are a snapshot and must be rechecked on official pages before purchase or deployment.\n- Cost examples are scenarios, not provider quotes.\n- Recommendations must pass the HisabDo-specific evaluation gates before real financial data is used.\n- Open the documents in Microsoft Word and insert/update a Table of Contents from the built-in Heading styles if required.\n'''
(OUT/'README.txt').write_text(readme,encoding='utf-8')

zip_path=Path('/home/user/HisabDo_AI_Model_API_Research_Package.zip')
with zipfile.ZipFile(zip_path,'w',zipfile.ZIP_DEFLATED) as z:
    for p in sorted(OUT.iterdir()): z.write(p,arcname=p.name)
print(zip_path)
for p in sorted(OUT.iterdir()): print(p.name,p.stat().st_size)
